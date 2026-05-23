from docx import Document
import io

def build_document(sector_name, sector_info):
    """Returns .docx file as bytes"""
    entities = sector_info.get("entities", {})
    inconsistencies = sector_info.get("inconsistencies", [])
    description = sector_info.get("description", "")

    doc = Document()

    # Title
    doc.add_heading(f"{sector_name} Capstone Project", level=0)
    doc.add_paragraph(f"Sector: {sector_name}  |  {description}")

    # Business Overview
    doc.add_heading("Business Overview", level=1)
    doc.add_paragraph(
        f"This project builds a complete enterprise-grade data engineering "
        f"and analytics solution for the {sector_name} sector.\n\n"
        "Modules covered: Data Ingestion, Python Automation, MongoDB CRUD, "
        "MySQL Stored Procedures, PySpark Transformations, Power BI Dashboards, "
        "Unix Shell Scripting."
    )

    # Entity Schemas
    doc.add_heading("Entity Schemas & Data Dictionary", level=1)
    for entity_name, entity_info in entities.items():
        doc.add_heading(entity_name, level=2)
        doc.add_paragraph(
            f"File: {entity_info.get('file','N/A')}  |  "
            f"Format: {entity_info.get('format','N/A')}  |  "
            f"Records: {entity_info.get('num_records','N/A')}"
        )
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Column"
        hdr[1].text = "Type"
        hdr[2].text = "Flag"
        hdr[3].text = "Description"
        for col in entity_info.get("columns", []):
            row = table.add_row().cells
            row[0].text = col[0]
            row[1].text = col[2]
            row[2].text = col[3]
            row[3].text = col[1]
        doc.add_paragraph("")

    # Inconsistencies
    if inconsistencies:
        doc.add_heading("Data Inconsistencies to Handle", level=1)
        for item in inconsistencies:
            doc.add_paragraph(item, style="List Bullet")

    # User Stories
    sections = {
        "Unix & Shell": [
            "Count total records in the primary CSV file.",
            "Filter records where billing amount exceeds 10,000.",
            "Identify and count duplicate IDs.",
            "Extract records with missing fields.",
            "Sort records by amount descending and display top 10.",
        ],
        "Python Data Cleaning": [
            "Remove duplicate rows keeping the first occurrence.",
            "Normalize date columns to YYYY-MM-DD format.",
            "Fill missing categorical values with 'Unknown'.",
            "Validate and zero-fill invalid numeric amount fields.",
            "Export cleaned data to a new CSV file.",
        ],
        "MongoDB": [
            "Insert all records from the cleaned CSV into a MongoDB collection.",
            "Update billing amount for a specific category.",
            "Set a default value where a field is null or empty.",
            "Query high-risk records by age and amount thresholds.",
            "Export specific fields to CSV using mongoexport.",
        ],
        "MySQL": [
            "Create a stored function to categorize billing as Low/Medium/High.",
            "Create a stored function to classify records by age group.",
            "Create a stored procedure to summarize by age and billing category.",
            "Use JOINs to combine entity tables for reporting.",
        ],
        "PySpark": [
            "Load CSV data into a Spark DataFrame with header inference.",
            "Drop fully null rows and trim whitespace from Name column.",
            "Join with location CSV and group by city to count records.",
            "Compute top 3 categories by average billing amount.",
            "Filter records admitted during Monsoon 2023 (Jul-Sep).",
        ],
        "Power BI / DAX": [
            "Create a DAX measure for total record count and billing.",
            "Implement Month-over-Month growth calculation.",
            "Rank locations by total billing using RANKX.",
            "Create an age bucket calculated column.",
            "Build a rolling 30-day average billing measure.",
        ],
    }

    doc.add_heading("User Stories", level=1)
    for section_title, stories in sections.items():
        doc.add_heading(section_title, level=2)
        for i, story in enumerate(stories, 1):
            doc.add_paragraph(
                f"US-{i:02d}: As a data engineer, I want to {story}",
                style="List Bullet"
            )

    # Save to bytes and return
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
